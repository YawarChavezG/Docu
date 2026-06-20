"""
Servicio: titulo_formulario_service
Extrae el titulo de formularios (.xlsx, .docx, .doc).
Busca el titulo en el contenido del archivo, no en el nombre del archivo.

Estrategia:
- .xlsx (Excel): lee celda C3 (fila 3, col C) del primer sheet.
  En los formularios COFAR, ahi esta el titulo real.
- .docx (Word): lee el primer parrafo relevante despues del codigo,
  similar a caratula_service.py.
- .doc (Word legacy): no se puede parsear facilmente. Fallback a nombre de archivo.
"""
import io
import logging
import re
from typing import Optional

logger = logging.getLogger(__name__)


def extraer_titulo_xlsx(contenido: bytes) -> Optional[str]:
    """
    Lee el titulo de un .xlsx desde el contenido del archivo.
    
    Estrategia: en los formularios COFAR, la estructura tipica es:
      Fila 1: area/departamento (ej: "PLANIFICACION INDUSTRIAL")
      Fila 2: "FORMULARIO: {codigo}" 
      Fila 3: TITULO REAL DEL DOCUMENTO (ej: "PROGRAMA DE EVALUACION...")
    
    Buscamos en filas 2-5 cualquier celda con texto largo (>15 chars)
    que NO contenga "FORMULARIO", "Versi", "Vigente", ni "NRO".
    Priorizamos la fila 3 sobre las demas.
    """
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(contenido), data_only=True)
        ws = wb.active
        if ws:
            skip_keywords = ['formulario', 'versi', 'vigente', 'nro de', 'aviso de', 'fecha de', 'datos del']
            candidates = []
            for row_num in range(1, min(ws.max_row + 1, 8)):
                for col_num in range(1, min(ws.max_column + 1, 10)):
                    val = ws.cell(row=row_num, column=col_num).value
                    if val and isinstance(val, str) and len(val.strip()) > 10:
                        txt = val.strip()
                        lower = txt.lower()
                        if not any(k in lower for k in skip_keywords):
                            candidates.append((row_num, txt))
            # Priorizar fila 3, luego 2, luego 4, luego 1
            for priority_row in [3, 2, 4, 1]:
                for r, txt in candidates:
                    if r == priority_row:
                        return txt[:200]
            if candidates:
                return candidates[0][1][:200]
    except Exception as e:
        logger.warning(f"Error extrayendo titulo de xlsx: {e}")
    return None


def extraer_titulo_docx(contenido: bytes) -> Optional[str]:
    """Lee el primer parrafo significativo del .docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(contenido))
        # Buscar en parrafos
        for para in doc.paragraphs:
            t = para.text.strip()
            if len(t) > 10 and not t.startswith('[') and 'pagina' not in t.lower():
                return t[:200]
        # Buscar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if len(t) > 10:
                        return t[:200]
    except Exception as e:
        logger.warning(f"Error extrayendo titulo de docx: {e}")
    return None


def extraer_titulo(contenido: bytes, filename: str) -> Optional[str]:
    """
    Extrae el titulo del formulario segun su tipo.
    Args:
        contenido: bytes del archivo
        filename: nombre original del archivo (para extension y fallback)
    Returns:
        Titulo extraido o None si no se pudo.
    """
    ext = (filename or '').lower()
    titulo = None

    if ext.endswith('.xlsx') or ext.endswith('.xls'):
        titulo = extraer_titulo_xlsx(contenido)
    elif ext.endswith('.docx'):
        titulo = extraer_titulo_docx(contenido)
    elif ext.endswith('.doc'):
        # .doc legacy: dificil de parsear, mejor fallback
        pass

    if titulo:
        return titulo

    # Fallback: extraer del nombre del archivo (quitar codigo y version)
    name_no_ext = re.sub(r'\.[^.]+$', '', filename or '')
    # Quitar patron de codigo al inicio (ej: "ACD-5-008-F01-F08 ")
    name_clean = re.sub(r'^[A-Z0-9]{2,6}[\s-]?\d[\d\s-]*[A-Z]*\d*[\s-]*', '', name_no_ext, count=1)
    # Quitar VXX del final
    name_clean = re.sub(r'\s*V\d{2}\s*$', '', name_clean).strip()
    if name_clean:
        return name_clean

    return filename or 'Formulario'
