"""
Servicio: caratula_service
R3 item 0.3: Validacion de caratula del documento .docx.

Lee la primera pagina del .docx (caratula) y extrae 3 campos:
- titulo (texto en MAYUSCULAS o normalizado)
- codigo (formato "CC-3-005")
- version (formato "00")

Segun PROPUESTA-R3-TABLAS.md §1.5.2:

    Validacion de caratula (importante):
    Cuando el usuario sube el archivo .docx (documento principal), el
    backend debe leer la primera pagina del .docx y verificar:
    - El titulo en la caratula coincide con documentos.titulo
    - El codigo coincide con documentos.codigo
    - La version coincide con version_snapshot

    Si NO coincide: advertencia (no bloqueante). Si coincide: confirmacion.

Estrategia de extraccion:
1. Leer las primeras N lineas de texto del documento (paragrafos, tablas).
2. Buscar patrones regex:
   - Codigo: [A-Z]{2,5}-d{1,2}-d{3,4} (sigla-tipo-correlativo)
   - Version: (V|VERSION)[s:]*0?(d{1,2}) (con o sin prefijo V)
   - Titulo: texto despues del codigo, antes de la version (heuristica).

La extraccion NO es 100% precisa (depende del formato del .docx de
cada area), pero es SUFICIENTE para una validacion warning-only.
Si falla la extraccion, el servicio retorna None en los campos y la
validacion se considera "no se pudo validar" (NO warning).
"""
import io
import logging
import re
from dataclasses import dataclass
from typing import Optional

logger = logging.getLogger(__name__)


# Regex patrones para extraer campos de la caratula.
# El codigo es "CC-3-005" (sigla 2-5 letras - tipo 1-2 digitos - correlativo 3-4 digitos).
_RE_CODIGO = re.compile(r"\b([A-Z]{2,5}-\d{1,2}-\d{3,4})\b")
# Version puede aparecer como "V00", "V 00", "VERSION: 00", "Version 0", etc.
_RE_VERSION = re.compile(r"(?i)\b(?:v|version|ver\.?)\s*[:.]?\s*0?(\d{1,2})\b")


@dataclass
class CaratulaExtraida:
    """Resultado de leer la caratula de un .docx."""
    codigo: Optional[str] = None        # "CC-3-005" o None si no se encontro
    version: Optional[str] = None       # "00" o None si no se encontro
    titulo: Optional[str] = None        # texto del titulo o None
    texto_crudo: str = ""               # primeras lineas del doc (para debug)
    exitoso: bool = False               # True si se pudo leer el doc


def _leer_primeras_lineas(docx_bytes: bytes, max_caracteres: int = 3000) -> str:
    """
    Lee las primeras lineas de texto del .docx.
    Incluye texto de paragrafos Y tablas (cubre el caso comun donde la
    caratula es una tabla con Codigo/Titulo/Version en celdas separadas).
    """
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx no instalado; validacion de caratula deshabilitada")
        return ""
    try:
        doc = Document(io.BytesIO(docx_bytes))
    except Exception as e:
        logger.warning(f"No se pudo parsear el .docx: {e}")
        return ""
    lineas = []
    chars = 0
    for para in doc.paragraphs:
        if not para.text:
            continue
        lineas.append(para.text)
        chars += len(para.text)
        if chars >= max_caracteres:
            break
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                t = celda.text.strip()
                if t:
                    lineas.append(t)
                    chars += len(t)
                    if chars >= max_caracteres:
                        break
            if chars >= max_caracteres:
                break
        if chars >= max_caracteres:
            break
    return "\n".join(lineas[:30])  # max 30 lineas


def _extraer_codigo(texto: str) -> Optional[str]:
    """Busca el primer codigo de documento en el texto."""
    m = _RE_CODIGO.search(texto)
    if m:
        return m.group(1)
    return None


def _extraer_version(texto: str) -> Optional[str]:
    """Busca la primera version en formato V00, V 01, etc."""
    m = _RE_VERSION.search(texto)
    if m:
        # Zero-pad a 2 digitos
        return f"{int(m.group(1)):02d}"
    return None


def _extraer_titulo(texto: str, codigo: Optional[str], version: Optional[str]) -> Optional[str]:
    """
    Heuristica: el titulo es el texto entre el codigo y la version.
    Si no se encuentran marcadores, retorna la linea mas larga
    (excluyendo lineas con el codigo o version).
    """
    lineas = [l.strip() for l in texto.split("\n") if l.strip()]
    if codigo is None and version is None:
        # Sin marcadores, retornar la linea mas larga que no sea solo numerica
        candidatas = [l for l in lineas if len(l) > 10 and not l.isdigit()]
        if candidatas:
            return max(candidatas, key=len)[:200]
        return None
    # Filtrar lineas con codigo o version
    candidatas = [
        l for l in lineas
        if (codigo is None or codigo not in l)
        and (version is None or not _RE_VERSION.search(l))
        and len(l) > 5
    ]
    if candidatas:
        # Tomar la primera candidata (suele estar entre codigo y version)
        return candidatas[0][:200]
    return None


def extraer_caratula(docx_bytes: bytes) -> CaratulaExtraida:
    """
    Lee la primera pagina del .docx y extrae codigo, version y titulo.

    Args:
        docx_bytes: contenido del .docx como bytes.

    Returns:
        CaratulaExtraida con los campos encontrados (None si no se
        encontro). El flag `exitoso` indica si el .docx se pudo parsear.
    """
    texto = _leer_primeras_lineas(docx_bytes)
    if not texto:
        return CaratulaExtraida(texto_crudo="", exitoso=False)

    codigo = _extraer_codigo(texto)
    version = _extraer_version(texto)
    titulo = _extraer_titulo(texto, codigo, version)

    return CaratulaExtraida(
        codigo=codigo,
        version=version,
        titulo=titulo,
        texto_crudo=texto[:500],  # solo 500 chars para no inflar el log
        exitoso=True,
    )


@dataclass
class ResultadoValidacion:
    """Resultado de comparar la caratula con los datos esperados."""
    coincide: bool
    warnings: list[str]        # mensajes de advertencia (NO bloqueantes)
    caratula: CaratulaExtraida

    def to_dict(self) -> dict:
        return {
            "coincide": self.coincide,
            "warnings": self.warnings,
            "caratula": {
                "codigo": self.caratula.codigo,
                "version": self.caratula.version,
                "titulo": self.caratula.titulo,
                "exitoso": self.caratula.exitoso,
            },
        }


def validar_caratula(
    docx_bytes: bytes,
    codigo_esperado: str,
    version_esperada: str,
    titulo_esperado: str,
) -> ResultadoValidacion:
    """
    Lee la caratula del .docx y compara con los datos esperados del documento.

    Si NO se puede parsear el .docx, no se encuentra codigo NI version
    (no hay marcadores de caratula), o los campos son ambiguos: retorna
    `coincide=True` y `warnings=[]` (no se puede validar de forma
    confiable, asi que NO penalizamos al usuario).

    Si se puede parsear pero hay diferencias: retorna `coincide=False`
    y `warnings=[...]` con las diferencias encontradas.

    Args:
        docx_bytes: contenido del .docx.
        codigo_esperado: codigo del documento (ej: "CC-3-005").
        version_esperada: version del documento (ej: "00").
        titulo_esperado: titulo del documento (case-insensitive match).

    Returns:
        ResultadoValidacion con detalles de la comparacion.
    """
    caratula = extraer_caratula(docx_bytes)
    if not caratula.exitoso:
        return ResultadoValidacion(
            coincide=True,
            warnings=[],
            caratula=caratula,
        )

    # Si no hay NI codigo NI version, no podemos validar de forma confiable
    # (el .docx probablemente no tiene una caratula estructurada).
    if caratula.codigo is None and caratula.version is None:
        return ResultadoValidacion(
            coincide=True,
            warnings=[],
            caratula=caratula,
        )

    warnings = []
    if caratula.codigo is not None and caratula.codigo != codigo_esperado:
        warnings.append(
            f"Codigo en caratula ({caratula.codigo}) no coincide con el esperado ({codigo_esperado})"
        )
    if caratula.version is not None and caratula.version != version_esperada:
        warnings.append(
            f"Version en caratula ({caratula.version}) no coincide con la esperada ({version_esperada})"
        )
    if caratula.titulo is not None and titulo_esperado:
        # Match case-insensitive, ignorando espacios multiples
        titulo_norm = " ".join(caratula.titulo.lower().split())
        esperado_norm = " ".join(titulo_esperado.lower().split())
        if titulo_norm != esperado_norm:
            # No es bloqueante: puede haber variaciones de formato
            warnings.append(
                f"Titulo en caratula ({caratula.titulo!r}) no coincide exactamente con el esperado "
                f"({titulo_esperado!r}). Verificar manualmente."
            )

    return ResultadoValidacion(
        coincide=len(warnings) == 0,
        warnings=warnings,
        caratula=caratula,
    )
