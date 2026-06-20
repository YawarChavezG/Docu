"""
Helpers para tests: crea un .docx en memoria con la caratula especificada.
Usado por tests del servicio de validacion de caratula (R3 item 0.3).
"""
from io import BytesIO

from docx import Document
from docx.shared import Pt


def crear_docx_caratula(
    codigo: str = "CC-3-005",
    titulo: str = "PROCEDIMIENTO DE MUESTREO",
    version: str = "00",
) -> bytes:
    """
    Crea un .docx en memoria con una caratula simple (3 lineas: codigo,
    titulo, version). Usado para tests del servicio de validacion.
    """
    doc = Document()
    p1 = doc.add_paragraph(codigo)
    p1.runs[0].font.size = Pt(14)
    p1.runs[0].bold = True
    p2 = doc.add_paragraph(titulo)
    p2.runs[0].font.size = Pt(18)
    p2.runs[0].bold = True
    p3 = doc.add_paragraph(f"V{version}")
    p3.runs[0].font.size = Pt(12)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def crear_docx_sin_caratula(texto_libre: str = "lorem ipsum dolor sit amet") -> bytes:
    """Crea un .docx sin codigo/titulo/version reconocibles."""
    doc = Document()
    doc.add_paragraph(texto_libre)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def crear_docx_con_header(header_text: str = "", body_text: str = "") -> bytes:
    """Crea un .docx con texto en el header (section.header)."""
    from docx.enum.section import WD_HEADER_FOOTER
    doc = Document()
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = header_text
    if body_text:
        doc.add_paragraph(body_text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def crear_docx_con_header_tabla(
    header_cells: list[str],
    body_text: str = "",
) -> bytes:
    """Crea un .docx con una tabla en el header (codigo/titulo/version en celdas)."""
    from docx.enum.section import WD_HEADER_FOOTER
    from docx import Document as DsDocument
    from docx.shared import Inches
    doc = DsDocument()
    section = doc.sections[0]
    header = section.header
    # Agregar tabla al header (width requerido en python-docx para header)
    table = header.add_table(rows=len(header_cells), cols=1, width=Inches(6))
    for i, text in enumerate(header_cells):
        table.cell(i, 0).text = text
    if body_text:
        doc.add_paragraph(body_text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
